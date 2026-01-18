from pathlib import Path

from docx import Document
from odf import teletype
from odf import text as odf_text
from odf.opendocument import load
from pypdf import PdfReader

from resume_ragbot.extraction.base import ExtractedDocument, Extractor

PLAINTEXT_EXTENSIONS = {".txt", ".text", ".md"}


class DocxExtractor(Extractor):
    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def extract(self, file_path: Path) -> ExtractedDocument:
        doc = Document(str(file_path))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        return ExtractedDocument(
            content=content,
            source=file_path.name,
            metadata={"file_type": file_path.suffix.lower()},
        )


class OdtExtractor(Extractor):
    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".odt"

    def extract(self, file_path: Path) -> ExtractedDocument:
        doc = load(str(file_path))

        paragraphs: list[str] = []
        for p in doc.getElementsByType(odf_text.P):
            text_content = teletype.extractText(p)
            if text_content.strip():
                paragraphs.append(text_content)

        content = "\n\n".join(paragraphs)

        return ExtractedDocument(
            content=content,
            source=file_path.name,
            metadata={"file_type": file_path.suffix.lower()},
        )


class PdfExtractor(Extractor):
    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def extract(self, file_path: Path) -> ExtractedDocument:
        reader = PdfReader(file_path)

        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text)

        content = "\n\n".join(pages)

        return ExtractedDocument(
            content=content,
            source=file_path.name,
            metadata={
                "file_type": file_path.suffix.lower(),
                "page_count": str(len(reader.pages)),
            },
        )


class PlaintextExtractor(Extractor):
    """Plaintext extractor. Assumes UTF-8 encoding."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in PLAINTEXT_EXTENSIONS

    def extract(self, file_path: Path) -> ExtractedDocument:
        content = file_path.read_text(encoding="utf-8")

        return ExtractedDocument(
            content=content,
            source=file_path.name,
            metadata={"file_type": file_path.suffix.lower()},
        )


EXTRACTORS = [
    DocxExtractor(),
    OdtExtractor(),
    PdfExtractor(),
    PlaintextExtractor(),
]
